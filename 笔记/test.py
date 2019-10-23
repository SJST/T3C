import uuid
import time
import os
import xlwt
import re
from ExeclCL import GetExeclDate
from Mongodb import Mongodb_client
import random
import json
import pymongo
import xlrd

#输入自定义字段
'''dlist=[{"a":1},{"b":2}]
for i in range(0,len(dlist)):
    conc=str(input("是否继续添加指定字段(Y/N):"))
    if conc=="Y":
        RelyDataKey=str(input("请输入字段名:"))
        RelyDataValue=str(input("请输入字段值:"))
        dlist.append({RelyDataKey:RelyDataValue})
    else:
        print("-------------自定义变量填写完成----------")
        break
print(dlist)'''
#赋值变量
'''a=1
b=a
print(b)
'''
#l两个字典合并 用一个字典的value 代替另一个字典的value
'''dict1={"a":1,"b":2,"c":3}
dict2={"a":"1","b":"2"}
dict1.update(dict2)
print(dict1)'''
#抛出异常
'''
try:
   
except Exception as e:
    print("获取swaggerdata失败,以下是失败信息" + "\n")
    print(e)
'''
#两个字典 比较
'''
dict1={"a":1,"b":2,"c":3}
dict2={"a":"1","b":"2"}
key11=dict1.keys()
key2=dict2.keys()
print(key11)
diff_data={}
for i in key11:
    if i not in key2:
        diff_data.update({str(i):dict1[i]})
print(diff_data)
'''

#dict 转str
'''
TestData={"a":1}
TestData.update({"b":1})
StrTestData=str(TestData)

print(StrTestData.replace(",",","+"\n"))
'''
#json dict 写入txt 不乱吗
'''

name = json.dumps('中国你好', ensure_ascii=False) #重点就是这一句代码 
date = time.strftime('%Y-%m-%d', time.localtime(time.time())) #此处是时间转换 
f = codecs.open(date + '.txt', 'a', 'utf-8') 
f.write('%s\n\n' % name) 
f.close()
'''

#将数据写入制定的文件中
'''
TestData={"a":1,"c":3}
TestData.update({"b":1})
# 时间转换
LocateTime=time.strftime('%Y-%m-%d %H:%M:%S',time.localtime(time.time()))
RQ=str(LocateTime).split()[0]
SJ=str(LocateTime).split()[1]
a="".join([ x for x in (RQ.split("-"))])
b="".join([ x for x in (SJ.split(":"))])
dictionaryMC=a+b
path="C:\\Users\\admin\\PycharmProjects\\TEST\\A_CaseData\\"+dictionaryMC
if not os.path.exists(path):
    os.mkdir(path)
f=path+"\\DiffSwaggerData.txt"
file=open(f,"w+",encoding="UTF-8")
file.write((str(TestData).replace(",",","+"\n")).replace("'",'"'))
file.close()
'''
'''创建目录
import os
# 创建的目录
path = "E:/ly/ly"
if not os.path.exists(path):
    os.mkdir(path)


'''
#遍历目录下的所有文件
'''
DIR="D:\\SJT\\docs\\20_工程过程\\30_设计相关\\30_概要设计\\各服务概要设计"
files_dir=[]
SMDDIR=[]
data=(os.walk(DIR))
for root, dirs, files in data:
    if re.search(r'\\SMD$',root):
        SMDDIR.append(root)
    else:pass

for item in SMDDIR:
    for root, dirs, files in os.walk(item):
        for i in files:
            Files=item+"\\"+("".join(i))
            files_dir.append(Files)
print(files_dir)
for  a in files_dir:
    try:
        SMDDataList=(GetExeclDate.ExeclData().GetExeclDate(a,"COL"))

        NeedList=GetExeclDate.ExeclData().NeedFeild(SMDDataList)
        print(NeedList)
        INSET=Mongodb_client.MongoCLient().MongoCLient("T3C","SMD_DATA","insert",NeedList)
    except Exception as e:
        print(e)
'''



'''
5、time.strftime的参数

strftime(format[, tuple]) -> string
将指定的struct_time(默认为当前时间)，根据指定的格式化字符串输出
python中时间日期格式化符号：
%y 两位数的年份表示（00-99）
%Y 四位数的年份表示（000-9999）
%m 月份（01-12）
%d 月内中的一天（0-31）
%H 24小时制小时数（0-23）
%I 12小时制小时数（01-12） 
%M 分钟数（00=59）
%S 秒（00-59）

%a 本地简化星期名称
%A 本地完整星期名称
%b 本地简化的月份名称
%B 本地完整的月份名称
%c 本地相应的日期表示和时间表示
%j 年内的一天（001-366）
%p 本地A.M.或P.M.的等价符
%U 一年中的星期数（00-53）星期天为星期的开始
%w 星期（0-6），星期天为星期的开始
%W 一年中的星期数（00-53）星期一为星期的开始
%x 本地相应的日期表示
%X 本地相应的时间表示
%Z 当前时区的名称
%% %号本身
'''
#str 定位元素 返回index
'''
strting="laay_old"
Index = strting.find("_")
New_Str=(strting.replace(strting[Index],""))
print(New_Str)
ResultStr=New_Str.replace(New_Str[Index],New_Str[Index].upper())
print(ResultStr)
'''
#将一个字典的值赋给相同key的另一个字典
'''
dict1={"a":"2400","b":"string","c":0}
dict2={"a":"法院","b":"VC600","c":0}

for i in dict2:
    print(i)
    if (dict1[i]) and (re.search(r'string',dict1[i])):
        dict1.update({i:dict2[i]})
print(dict1)
'''

#生成随机数字
'''
r=random.randint(0,20)
print(r)
'''
#字符串转float
#print(type(float("123.45")))
#生成uuid

# uid = str(uuid.uuid4())
# suid = ''.join(uid.split('-'))
# print(suid)

'''
    import uuid
     
    print uuid.uuid1()
    #bf1dfacf-67d8-11e8-9a23-408d5c985711
    print uuid.uuid3(uuid.NAMESPACE_DNS, 'yuanlin')
    #ddb366f5-d4bc-3a20-ac68-e13c0560058f
    print uuid.uuid4()
    #144d622b-e83a-40ea-8ca1-66af8a86261c
    print uuid.uuid5(uuid.NAMESPACE_DNS, 'yuanlin')
    #4a47c18d-037a-5df6-9e12-20b643c334d3

        乍一看全都是36个字符，那么他们到底有什么不同呢，下面一一分析。

        uuid1()：这个是根据当前的时间戳和MAC地址生成的，最后的12个字符408d5c985711对应的就是MAC地址，因为是MAC地址，那么唯一性应该不用说了。但是生成后暴露了MAC地址这就很不好了。

        uuid3()：里面的namespace和具体的字符串都是我们指定的，然后呢···应该是通过MD5生成的，这个我们也很少用到，莫名其妙的感觉。

        uuid4()：这是基于随机数的uuid，既然是随机就有可能真的遇到相同的，但这就像中奖似的，几率超小，因为是随机而且使用还方便，所以使用这个的还是比较多的。

        uuid5()：这个看起来和uuid3()貌似并没有什么不同，写法一样，也是由用户来指定namespace和字符串，不过这里用的散列并不是MD5，而是SHA1.

'''
#重新赋值
'''
a=["1","2"]
a="".join(a)
print(a)
'''
a=["1","2"]
print(a*2)
#Python操作docx
import docx

from docx import Document


path = "C:\\Users\\Administrator\\Desktop\\审判管理系统升级——数据库设计说明书 2.docx"
Name="审判管理系统升级——数据库设计说明书 2"
document = Document(path)
title_list=[]
title_list.append("项目")
title_list.append("表清单")
for paragraph in document.paragraphs:
    Par=paragraph.text
    if re.search(r'^表 ',Par) and not (re.search(r'列清单$',Par)):
        res_Par=str(str(Par).replace("表 ","")).replace("/","_")
        title_list.append(res_Par)
# print(len(title_list))
# print(title_list)

tables = document.tables
# print(len(tables))
# print(tables)
for table in tables:
    LieList=[]
    ja_list = []
    index_flags= tables.index(table)
    if index_flags not in [0,1]:
        len_row = len(table.rows)
        col_count = len(table.columns)
        for i in range(col_count):
            mubiao = table.cell(0, i).text
            for j in range(len_row):
                ja_dict = {}
                result_text = table.cell(j, i).text
                if result_text!=mubiao:
                    ja_dict.update({mubiao:result_text})
                    ja_list.append(ja_dict)
    if ja_list!=[]:
        frist_data=ja_list
        if frist_data:
            #for item in frist_data:

            # frist_dict={}
            # print(frist_data)
            # frist_dict.update({title_list[index_flags]:frist_data})
            # 创建一个workbook 设置编码
            workbook = xlwt.Workbook(encoding='utf-8')
            # 创建一个worksheet
            worksheet = workbook.add_sheet(title_list[index_flags],cell_overwrite_ok=True)
            ExeclTitle=[]
            all_nr=[]
            for i in frist_data:
                One_title=list(i.keys())[0]
                if One_title not in ExeclTitle:
                    ExeclTitle.append(One_title)
            for b in ExeclTitle:
                num=ExeclTitle.index(b)
                h_list=[]
                for y in frist_data:
                    key=list(y.keys())[0]

                    if key==b:
                        h_list.append(y.get(key))
                        if h_list not in LieList:
                            LieList.append(h_list)


            print(LieList)
            for a in range(0,len(ExeclTitle)):
                worksheet.write(0,a,ExeclTitle[a])
            for ChildList in LieList:
                #列数
                j=LieList.index(ChildList)
                for l in ChildList:
                    k=ChildList.index(l)
                    worksheet.write(k+1,j,l)


            workbook.save(Name+'.xls')





    if  tables.index(table)==2:
        break






